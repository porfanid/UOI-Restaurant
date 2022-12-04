from docx import Document
import re
from urllib.request import urlopen
from io import BytesIO

from datetime import datetime
currentMonth = datetime.now().month

day=1


def get_table_as_string(data):
	data_dict = {}
	for j in data.keys():
		if not j:
			continue
		data_dict[j]= [re.sub(' +', ' ', i.strip().replace("\n", " ")) for i in data[j].replace("ή", "").split("\n\n") if i]
	return data_dict

#Each week is in a seperate page. So, all this function does is to read each page and parse it's contents. This is the basic function of the program as it is responsible for getting the contents of the menu
def get_each_week(table):
	global day
	global currentMonth
	menu = {}
	# Data will be a list of rows represented as dictionaries
	# containing each row's data.
	data = []
	for i, row in enumerate(table.rows):
		text = (cell.text for cell in row.cells)

		# Establish the mapping based on the first row
		# headers; these will become the keys of our dictionary
		if i == 0:
			keys = tuple(text)
			continue

		# Construct a dictionary for this row, mapping
		# keys to values for this row
		row_data = dict(zip(keys, text))
		data.append(row_data)
	lunch_data = get_table_as_string(data[0])
	dinner_data = get_table_as_string(data[1])
	for key in lunch_data.keys():
		full_date = "{} {}/{}".format(key,day, currentMonth)
		menu[full_date]={}
		menu[full_date]["lunch"] = lunch_data[key]
		menu[full_date]["dinner"] = dinner_data[key]
		day+=1
	return menu

def get_full_menu(document):
	full_menu = {}
	for table in document.tables:
		weekly_menu = get_each_week(table)
		full_menu.update(weekly_menu)
	return full_menu



def get_file_from_url(url):
	file = urlopen(url).read()
	file = BytesIO(file)
	document = Document(file)
	table = document.tables[0]
	return get_full_menu(document)
